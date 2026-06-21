"""
Export Service — Generates PDF, DOCX, TXT, SRT, JSON, CSV, ZIP exports.
All formats use free/open-source libraries only.
"""
import os, uuid, json, csv, zipfile
from datetime import datetime
from typing import List
from loguru import logger
from sqlalchemy.orm import Session

from app.core.database import SessionLocal, ExportJob, Project, Transcript, Summary, Frame, ChatMessage, StudyAsset
from app.core.config import settings


class ExportService:

    @staticmethod
    def generate(export_id: str, project_id: str, fmt: str, content_types: List[str]):
        db = SessionLocal()
        try:
            # Update status
            db.query(ExportJob).filter(ExportJob.id == export_id).update({"status": "processing"})
            db.commit()

            out_dir = os.path.join(settings.export_dir, project_id)
            os.makedirs(out_dir, exist_ok=True)

            project = db.query(Project).filter(Project.id == project_id).first()
            transcript = db.query(Transcript).filter(Transcript.project_id == project_id).first()
            summaries = db.query(Summary).filter(Summary.project_id == project_id).all()
            frames = db.query(Frame).filter(Frame.project_id == project_id).order_by(Frame.timestamp_seconds).all()
            chat_msgs = db.query(ChatMessage).filter(ChatMessage.project_id == project_id).order_by(ChatMessage.created_at).all()
            study_assets = db.query(StudyAsset).filter(StudyAsset.project_id == project_id).all()

            data = {
                "project": project, "transcript": transcript,
                "summaries": summaries, "frames": frames,
                "chat_msgs": chat_msgs, "study_assets": study_assets,
                "content_types": content_types,
            }

            dispatch = {
                "pdf": ExportService._pdf,
                "docx": ExportService._docx,
                "txt": ExportService._txt,
                "json": ExportService._json,
                "srt": ExportService._srt,
                "vtt": ExportService._vtt,
                "csv": ExportService._csv,
                "zip": ExportService._zip,
            }

            fn = dispatch.get(fmt)
            if not fn:
                raise ValueError(f"Unsupported format: {fmt}")

            file_path = fn(out_dir, project_id, data)

            db.query(ExportJob).filter(ExportJob.id == export_id).update({
                "status": "complete", "file_path": file_path,
            })
            db.commit()
            logger.info(f"Export complete: {file_path}")

        except Exception as e:
            logger.error(f"Export failed: {e}")
            db.query(ExportJob).filter(ExportJob.id == export_id).update({
                "status": "failed",
            })
            db.commit()
        finally:
            db.close()

    # ── PDF ──────────────────────────────────────────────────────────────

    @staticmethod
    def _pdf(out_dir: str, project_id: str, data: dict) -> str:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            HRFlowable, PageBreak,
        )
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        path = os.path.join(out_dir, f"report_{project_id[:8]}.pdf")
        doc = SimpleDocTemplate(path, pagesize=A4, topMargin=20*mm, bottomMargin=20*mm)
        styles = getSampleStyleSheet()
        story = []

        # ── Title page
        title_style = ParagraphStyle("T", parent=styles["Title"], fontSize=24,
                                     textColor=colors.HexColor("#1e3a8a"), spaceAfter=8)
        sub_style = ParagraphStyle("S", parent=styles["Normal"], fontSize=11,
                                   textColor=colors.HexColor("#64748b"), spaceAfter=4)
        h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14,
                             textColor=colors.HexColor("#1e3a8a"), spaceBefore=12, spaceAfter=6)
        body = ParagraphStyle("B", parent=styles["Normal"], fontSize=10, leading=15,
                              textColor=colors.HexColor("#1e293b"))

        project = data["project"]
        title = project.title or project.source_url or project.source_filename or "Video Analysis"

        story.append(Paragraph("📹 VideoLM — AI Video Analysis Report", title_style))
        story.append(Paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", sub_style))
        story.append(Paragraph(f"Source: {title}", sub_style))
        if project.duration_seconds:
            m, s = int(project.duration_seconds // 60), int(project.duration_seconds % 60)
            story.append(Paragraph(f"Duration: {m}:{s:02d}", sub_style))
        story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#e2e8f0")))
        story.append(Spacer(1, 6*mm))

        # ── Summary
        summaries = [s for s in data["summaries"] if s.summary_type == "medium"]
        if summaries:
            story.append(Paragraph("AI Summary", h2))
            story.append(Paragraph(summaries[0].content, body))
            story.append(Spacer(1, 4*mm))

        # ── Transcript
        transcript = data["transcript"]
        if transcript and "transcript" in data["content_types"]:
            story.append(PageBreak())
            story.append(Paragraph("Full Transcript", h2))
            # Segment by segment
            for seg in (transcript.segments or [])[:100]:
                ts = seg.get("start", 0)
                m2, s2 = int(ts // 60), int(ts % 60)
                label = f"{m2:02d}:{s2:02d}"
                story.append(Paragraph(
                    f'<font color="#2563eb">[{label}]</font> {seg.get("text", "")}',
                    ParagraphStyle("seg", parent=body, fontSize=9, leading=13)
                ))

        # ── Key frames
        if "frames" in data["content_types"] and data["frames"]:
            story.append(PageBreak())
            story.append(Paragraph("Key Frames & Captions", h2))
            for frame in data["frames"][:20]:
                if frame.caption:
                    story.append(Paragraph(
                        f'<b>[{frame.timestamp_label}]</b> {frame.caption}', body
                    ))
                    if frame.ocr_text:
                        story.append(Paragraph(
                            f'OCR: {frame.ocr_text[:200]}',
                            ParagraphStyle("ocr", parent=body, fontSize=8,
                                           textColor=colors.HexColor("#64748b"), leftIndent=10)
                        ))
                    story.append(Spacer(1, 2*mm))

        # ── Study notes
        if "flashcards" in data["content_types"]:
            flash_assets = [a for a in data["study_assets"] if a.asset_type == "flashcard"]
            if flash_assets:
                story.append(PageBreak())
                story.append(Paragraph("Flashcards", h2))
                cards = flash_assets[0].content.get("flashcards", [])
                table_data = [["Question", "Answer"]]
                for card in cards[:20]:
                    table_data.append([card.get("question", ""), card.get("answer", "")])
                t = Table(table_data, colWidths=[110*mm, 60*mm])
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e3a8a")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                ]))
                story.append(t)

        doc.build(story)
        return path

    # ── DOCX ─────────────────────────────────────────────────────────────

    @staticmethod
    def _docx(out_dir: str, project_id: str, data: dict) -> str:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        project = data["project"]

        # Title
        title = doc.add_heading("VideoLM — AI Video Analysis", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        src = project.source_url or project.source_filename or "Video"
        doc.add_paragraph(f"Source: {src}")
        doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
        doc.add_paragraph()

        # Summary
        summaries = [s for s in data["summaries"] if s.summary_type == "medium"]
        if summaries:
            doc.add_heading("AI Summary", 1)
            doc.add_paragraph(summaries[0].content)

        # Bullet summary
        bullets = [s for s in data["summaries"] if s.summary_type == "bullets"]
        if bullets:
            doc.add_heading("Key Points", 1)
            for line in bullets[0].content.splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip().lstrip("•").strip(), style="List Bullet")

        # Transcript
        if "transcript" in data["content_types"] and data["transcript"]:
            doc.add_page_break()
            doc.add_heading("Full Transcript", 1)
            t = data["transcript"]
            for seg in (t.segments or [])[:200]:
                ts = seg.get("start", 0)
                m, s = int(ts // 60), int(ts % 60)
                p = doc.add_paragraph()
                run = p.add_run(f"[{m:02d}:{s:02d}] ")
                run.bold = True
                run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)
                p.add_run(seg.get("text", ""))

        path = os.path.join(out_dir, f"report_{project_id[:8]}.docx")
        doc.save(path)
        return path

    # ── TXT ──────────────────────────────────────────────────────────────

    @staticmethod
    def _txt(out_dir: str, project_id: str, data: dict) -> str:
        path = os.path.join(out_dir, f"transcript_{project_id[:8]}.txt")
        t = data["transcript"]
        with open(path, "w", encoding="utf-8") as f:
            f.write("VideoLM — Transcript\n")
            f.write("=" * 50 + "\n\n")
            if t:
                for seg in (t.segments or []):
                    ts = seg.get("start", 0)
                    m, s = int(ts // 60), int(ts % 60)
                    f.write(f"[{m:02d}:{s:02d}] {seg.get('text', '').strip()}\n")
        return path

    # ── JSON ─────────────────────────────────────────────────────────────

    @staticmethod
    def _json(out_dir: str, project_id: str, data: dict) -> str:
        path = os.path.join(out_dir, f"export_{project_id[:8]}.json")
        project = data["project"]
        t = data["transcript"]
        output = {
            "generated_at": datetime.utcnow().isoformat(),
            "project": {
                "id": project.id, "title": project.title,
                "source": project.source_url or project.source_filename,
                "duration": project.duration_seconds,
                "language": project.language,
            },
            "transcript": {
                "text": t.full_text if t else "",
                "language": t.language if t else "en",
                "segments": t.segments if t else [],
            } if t else None,
            "summaries": [
                {"type": s.summary_type, "model": s.model_used, "content": s.content}
                for s in data["summaries"]
            ],
            "frames": [
                {
                    "timestamp": f.timestamp_label,
                    "caption": f.caption,
                    "ocr_text": f.ocr_text,
                }
                for f in data["frames"]
            ],
        }
        with open(path, "w", encoding="utf-8") as f:
            json.dump(output, f, indent=2, ensure_ascii=False)
        return path

    # ── SRT ──────────────────────────────────────────────────────────────

    @staticmethod
    def _srt(out_dir: str, project_id: str, data: dict) -> str:
        path = os.path.join(out_dir, f"subtitles_{project_id[:8]}.srt")
        t = data["transcript"]
        with open(path, "w", encoding="utf-8") as f:
            for i, seg in enumerate((t.segments if t else []), 1):
                start = _srt_time(seg.get("start", 0))
                end = _srt_time(seg.get("end", seg.get("start", 0) + 3))
                f.write(f"{i}\n{start} --> {end}\n{seg.get('text', '').strip()}\n\n")
        return path

    # ── VTT ──────────────────────────────────────────────────────────────

    @staticmethod
    def _vtt(out_dir: str, project_id: str, data: dict) -> str:
        path = os.path.join(out_dir, f"subtitles_{project_id[:8]}.vtt")
        t = data["transcript"]
        with open(path, "w", encoding="utf-8") as f:
            f.write("WEBVTT\n\n")
            for i, seg in enumerate((t.segments if t else []), 1):
                start = _vtt_time(seg.get("start", 0))
                end = _vtt_time(seg.get("end", seg.get("start", 0) + 3))
                f.write(f"{i}\n{start} --> {end}\n{seg.get('text', '').strip()}\n\n")
        return path

    # ── CSV ──────────────────────────────────────────────────────────────

    @staticmethod
    def _csv(out_dir: str, project_id: str, data: dict) -> str:
        path = os.path.join(out_dir, f"transcript_{project_id[:8]}.csv")
        t = data["transcript"]
        with open(path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["index", "start", "end", "text", "speaker"])
            for i, seg in enumerate((t.segments if t else []), 1):
                writer.writerow([
                    i, seg.get("start", ""), seg.get("end", ""),
                    seg.get("text", ""), seg.get("speaker", ""),
                ])
        return path

    # ── ZIP bundle ───────────────────────────────────────────────────────

    @staticmethod
    def _zip(out_dir: str, project_id: str, data: dict) -> str:
        # Generate all sub-formats first
        files = []
        for fn, fmt in [
            (ExportService._pdf, "pdf"),
            (ExportService._docx, "docx"),
            (ExportService._txt, "txt"),
            (ExportService._json, "json"),
            (ExportService._srt, "srt"),
        ]:
            try:
                p = fn(out_dir, project_id, data)
                if p and os.path.exists(p):
                    files.append(p)
            except Exception as e:
                logger.warning(f"Sub-export {fmt} failed: {e}")

        zip_path = os.path.join(out_dir, f"videolm_bundle_{project_id[:8]}.zip")
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for fp in files:
                zf.write(fp, os.path.basename(fp))
        return zip_path


def _srt_time(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _vtt_time(seconds: float) -> str:
    return _srt_time(seconds).replace(",", ".")
